from Bancodedadosreg import * 
import warnings
warnings.filterwarnings('ignore')
import json
from docx import Document,shared
from numpy.random import seed
import pandas as pd
from sklearn.model_selection import train_test_split
import matplotlib.pyplot as plt
from sklearn.metrics import r2_score
import numpy as np
from sklearn import metrics
val_loss=[]
import joblib
from scipy.stats import uniform, truncnorm, randint
from sklearn.model_selection import RandomizedSearchCV
from sklearn.preprocessing import PolynomialFeatures
from sklearn.linear_model import LinearRegression
from sklearn.pipeline import make_pipeline

initDB() 
    

#        import tensorflow as tf   
 #       tf.compat.v1.logging.set_verbosity(tf.compat.v1.logging.ERROR)
  #      from tensorflow.keras.models import Sequential, load_model
   #     from tensorflow.keras.layers import Dense
    #    tf.compat.v1.set_random_seed(2)
     #   seed(1)
      #  from keras.wrappers.scikit_learn import KerasRegressor



def graph_scatter(pred,real,title,diretorio,saida,label):
    
    fig,ax =plt.subplots(dpi=200)  
    ax.scatter(real,pred,c='r',label=label)
    ax.legend()
    ax.axis('square')
    menor=min([ax.get_xlim()[0],ax.get_ylim()[0]])
    maior=max([ax.get_xlim()[1],ax.get_ylim()[1]])
    ax.set_xlim(menor,maior)
    ax.set_ylim(menor,maior)
    ax.plot([menor,maior], [menor,maior],c='black',linewidth=0.65)
    ax.title.set_text(title)
    ax.set(xlabel="Valor real ({})".format(saida),ylabel="Previsão ({})".format(saida))

    plt.savefig(diretorio,bbox_inches='tight')
    plt.clf()
    plt.close

            
def norm_x(x,train_stats_x):
  return (x - train_stats_x['mean'].values) / (train_stats_x['std'].values)

def norm_rev_x(x,train_stats_x):
  return ((x* (train_stats_x['std'].values))+train_stats_x['mean'].values)

def norm_y(x,train_stats_y):
  return (x - train_stats_y['mean']) / (train_stats_y['std'])

def norm_rev_y(x,train_stats_y):
  return ((x* (train_stats_y['std']))+train_stats_y['mean'])

class modelo():
    def __init__(self,referencia,tipo):
        initDB()
        self.referencia=referencia
        self.tipo=tipo
        self.name=referencia+'_'+tipo
        a=view_one(self.referencia,self.tipo)
        self.grandeza=a[0][1]
        self.ferramenta=a[0][3]
        self.material=a[0][4]
        self.numero=a[0][5]
        self.observacoes=a[0][6]
        self.u_velocidade=a[0][7]
        self.u_avanco=a[0][8]
        self.u_profundidade=a[0][9]
        self.u_saida=a[0][10]
        
        df = pd.read_csv('Arquivos\Dados/'+self.name+'_dados.csv',sep=';')
        self.df = df.sort_values(by=self.grandeza)
        x=self.df
        y=x.pop(self.grandeza)
        
        train_stats_x = x.describe()
        self.train_stats_x = train_stats_x.transpose()
        train_stats_y = y.describe()
        self.train_stats_y = train_stats_y.transpose()
        
        x=norm_x(x,self.train_stats_x).values
        y=norm_y(y,self.train_stats_y).values
        
        self.x_train, self.x_test, self.y_train, self.y_test = train_test_split(x, y, test_size=0.2,random_state=0)
    
        self.y_train_r=norm_rev_y(self.y_train,self.train_stats_y)
        self.y_test_r=norm_rev_y(self.y_test,self.train_stats_y)
        

        self.treino_x=l=pd.DataFrame(self.x_train)
        self.treino_x=norm_rev_x(self.treino_x,self.train_stats_x)
        self.treino_x.columns=['n','f','a']
        
        self.treino_y=l=pd.DataFrame(self.y_train)
        self.treino_y=norm_rev_y(self.treino_y,self.train_stats_y)
        self.treino_y.columns=[self.grandeza]
        
        self.treino=np.round(self.treino_y.join(self.treino_x),2)
        
        self.teste_x=l=pd.DataFrame(self.x_test)
        self.teste_x=norm_rev_x(self.teste_x,self.train_stats_x)
        self.teste_x.columns=['n','f','a']
    
        self.teste_y=l=pd.DataFrame(self.y_test)
        self.teste_y=norm_rev_y(self.teste_y,self.train_stats_y)
        self.teste_y.columns=[self.grandeza]
    
        self.teste=np.round(self.teste_y.join(self.teste_x),2)


      
    def criar_modelos(self,cv,ite):  
        
        def tabela(df,doc):
            t = doc.add_table(df.shape[0]+1, df.shape[1])
            t.style = 'Medium List 1 Accent 1'
            
            
            for j in range(df.shape[-1]):
                t.cell(0,j).text = df.columns[j]
                
            for i in range(df.shape[0]):
                for j in range(df.shape[-1]):
                    t.cell(i+1,j).text = str(df.values[i,j])
        
            for row in t.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    paragraph = paragraphs[0]
                    run_obj = paragraph.runs
                    run = run_obj[0]
                    font = run.font
                    font.size =  shared.Pt(11)
                    
        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = shared.Pt(12)
        
        a=document.add_heading('Informações do estudo')
        a.alignment=1
        document.add_paragraph('Referência: {}'.format(self.referencia))
        document.add_paragraph('Grandeza: {}'.format(self.grandeza))
        document.add_paragraph('Tipo: {}'.format(self.tipo))
        document.add_paragraph('Material: {}'.format(self.material))
        document.add_paragraph('Ferramenta: {}'.format(self.ferramenta))
        document.add_paragraph('Número de experimentos: {}'.format(self.numero))
        document.add_paragraph('Observações:\n{}'.format(self.observacoes))
        document.add_heading('Unidades')
        document.add_paragraph('Velocidade: {}'.format(self.u_velocidade))
        document.add_paragraph('Avanço: {}'.format(self.u_avanco))
        document.add_paragraph('Profundidade de corte: {}'.format(self.u_profundidade))
        document.add_paragraph('{}: {}'.format(self.grandeza,self.u_saida))
        document.add_heading('Dados de teste')
        tabela(self.teste,document)
        document.add_heading('Dados de treino')
        tabela(self.treino,document)
        document.add_page_break()
        
        self.dici_erros=[]
        self.data_train_erro={}
        self.data_test_erro={}
        def erro(y_test_r,pred_test,y_train_r,pred_train,diretorio,name):
            
            erro_test=np.array([])
            erro_train=np.array([])
            
            for i in range(len(y_train_r)):
                erro_train=np.append(erro_train,abs(pred_train[i]-y_train_r[i])*100/y_train_r[i])
                
            for i in range(len(y_test_r)):
                erro_test=np.append(erro_test,abs(pred_test[i]-y_test_r[i])*100/y_test_r[i])
                
                
            erro_dic_train={diretorio+' Previsto': pred_train, diretorio+' Erro (%)': erro_train}
            erro_dic_test={diretorio+' Previsto': pred_test, diretorio+' Erro (%)': erro_test}
            
            self.data_train_erro.update(erro_dic_train)
            self.data_test_erro.update(erro_dic_test)
            
            
            errorel_test=np.round(sum(erro_test)/len(erro_test),2)
            errorel_train=np.round(sum(erro_train)/len(erro_train),2)
        
            
            cor_train_full=np.corrcoef(y_train_r,pred_train)[1,0]
            cor_test_full=np.corrcoef(y_test_r,pred_test)[1,0]
            cor_train=np.round(np.corrcoef(y_train_r,pred_train)[1,0],2)
            cor_test=np.round(np.corrcoef(y_test_r,pred_test)[1,0],2)
            det_test = np.round(cor_test_full**2,2)
            det_train = np.round(cor_train_full**2,2)
            mse_test = np.round(metrics.mean_squared_error(pred_test,y_test_r),2)
            mse_train = np.round(metrics.mean_squared_error(pred_train,y_train_r),2)
            rmse_test = np.round(np.sqrt(mse_test),2)
            rmse_train = np.round(np.sqrt(mse_train),2)
            
            self.dici_erros={'errorel_test': errorel_test, 'errorel_train': errorel_train,
                        'cor_test': cor_test, 'cor_train': cor_train,
                        'det_test': det_test, 'det_train': det_train,
                        'mse_test': mse_test , 'mse_train': mse_train,
                        'rmse_test': rmse_test, 'rmse_train': rmse_train}
            
            with open('Arquivos/'+diretorio+'/Erros/'+name+'_erros.json', 'w') as json_file:
                json.dump(self.dici_erros, json_file, indent=4)
        
        fig,ax =plt.subplots(dpi=200)
        fig2,ax2 =plt.subplots(dpi=200)
        
        
                
        def graficos_ind(pred_test,y_test_r,pred_train,y_train_r,u_saida, diretorio,name,label_test,label_train):
            graph_scatter(pred_test,y_test_r,'Previsão (dados teste) - '+diretorio,
                          'Arquivos/'+diretorio+'/Gráficos/'+name+'_teste.png',u_saida,label_test)
            graph_scatter(pred_train,y_train_r,'Previsão (dados treino) - '+diretorio,
                          'Arquivos/'+diretorio+'/Gráficos/'+name+'_treino.png',u_saida,label_train)
        
        ##################### Criação rede neural ################
        
   
        import tensorflow as tf   
        tf.compat.v1.logging.set_verbosity(tf.compat.v1.logging.ERROR)
        from tensorflow.keras.models import Sequential, load_model
        from tensorflow.keras.layers import Dense
        tf.compat.v1.set_random_seed(2)
        seed(1)
        from keras.wrappers.scikit_learn import KerasRegressor

                
        
        def create_model(n,ln,fun,layer2,epochs):
            
            initializer = tf.compat.v1.keras.initializers.glorot_uniform(seed=0)
            model = Sequential()
            model.add(Dense(int(n), input_dim=self.x_train.shape[1], activation=fun,
                            use_bias=True,kernel_initializer=initializer))
            if layer2==True:
                model.add(Dense(n, activation=fun,use_bias=True,kernel_initializer=initializer)) # Hidden 2
            model.add(Dense(1,kernel_initializer=initializer)) 
            optimizer = tf.keras.optimizers.Adam(learning_rate=ln)
            model.compile(loss='mean_squared_error', optimizer=optimizer,metrics=['mse'])
            return model
    
        ln = [0.1,0.01,0.001, 0.0001]
        n=randint(1,100)
        layer2=[False,True]
        epochs=randint(1,1000)
        fun=['tanh','relu']
        
        param_grid = dict(n=n,ln=ln,fun=fun,layer2=layer2,epochs=epochs)
    
        model = KerasRegressor(build_fn=create_model, verbose=0)

        random = RandomizedSearchCV(estimator=model, param_distributions=param_grid, n_jobs=1, cv=cv,n_iter=ite,verbose=2)

        random_result =random.fit(self.x_train, self.y_train)
        
        arquivo=open('Arquivos/RN/Iterações/'+self.name+'_iteraçoes.txt','w')
        
        arquivo.write("Melhores parâmetros: %f com %s" % (random_result.best_score_, random_result.best_params_))
        arquivo.write('\n\n')
        
        means = random_result.cv_results_['mean_test_score']
        stds = random_result.cv_results_['std_test_score']
        params = random_result.cv_results_['params']
        
        iteracao={'Média': [],'Desvio': [],'n': [], 'ln': [],'2° camada': [], 'Função': [],'Épocas': []}
        
        for mean, stdev, param in zip(means, stds, params):
            
            iteracao['Média'].append(round(mean,4))
            iteracao['Desvio'].append(round(stdev,4))
            iteracao['n'].append(param['n'])
            iteracao['ln'].append(param['ln'])
            iteracao['2° camada'].append(param['layer2'])
            iteracao['Função'].append(param['fun'])
            iteracao['Épocas'].append(param['epochs'])
        
        self.iteracao=pd.DataFrame.from_dict(iteracao)
        arquivo.write(self.iteracao.to_string())
        arquivo.close()
        
        with open('Arquivos/RN/Parâmetros/'+self.name+'_parametros.json', 'w') as json_file:
            json.dump(random_result.best_params_, json_file, indent=4)
            
        self.model=create_model(**random_result.best_params_) 
        history=self.model.fit(self.x_train,self.y_train,
                          validation_data=(self.x_test,self.y_test),
                          verbose=0,epochs=random_result.best_params_['epochs'])
        history
        self.model.save('Arquivos/RN/Modelos/'+self.name+'_modelo.mdl')
        
        pred_train0=self.model.predict(self.x_train)
        pred_train0_r=norm_rev_y(pred_train0,self.train_stats_y)
        pred_train=np.array([])
        for item in pred_train0_r:
            pred_train=np.append(pred_train,round(float(item),2))
        
        pred_test0 = self.model.predict(self.x_test)
        pred_test0_r=norm_rev_y(pred_test0,self.train_stats_y)
        pred_test=np.array([])
        for item in pred_test0_r:
            pred_test=np.append(pred_test,round(float(item),2))
            
        cor_train_full=np.corrcoef(self.y_train_r,pred_train)[1,0]
        cor_test_full=np.corrcoef(self.y_test_r,pred_test)[1,0]
        
        det_test = np.round(cor_test_full**2,2)
        det_train = np.round(cor_train_full**2,2)
        
        erro(self.y_test_r,pred_test,self.y_train_r,pred_train,'RN',self.name)
        graficos_ind(pred_test,self.y_test_r,pred_train,self.y_train_r,
                     self.u_saida, 'RN',self.name,'RN $R^2$={:.2f}'.format(det_test),
                     'RN $R^2$={:.2f}'.format(det_train))
        
        cor_train_full=np.corrcoef(self.y_train_r,pred_train)[1,0]
        cor_test_full=np.corrcoef(self.y_test_r,pred_test)[1,0]
        
        det_test = np.round(cor_test_full**2,2)
        det_train = np.round(cor_train_full**2,2)
        
        ax.scatter(self.y_test_r,pred_test,label='RN $R^2$={:.2f}'.format(det_test),marker='s')
        ax2.scatter(self.y_train_r,pred_train,label='RN $R^2$={:.2f}'.format(det_train),marker='s')
        
        arquivo2=open('Arquivos/RN/Pesos/'+self.name+'_pesos.txt','w')
        k=self.model.get_weights()

        arquivo2.write('Pesos - camada oculta 1')
        arquivo2.write('\n')
        arquivo2.write(str(k[0]))
        arquivo2.write('\n')
        arquivo2.write('Bias - camada oculta')
        arquivo2.write('\n')
        arquivo2.write(str(k[1]))
        arquivo2.write('\n')
        if random_result.best_params_['layer2']==False:
            arquivo2.write('Pesos - camada saída')
            arquivo2.write('\n')
            arquivo2.write(str(k[2].transpose()))
            arquivo2.write('\n')

        elif random_result.best_params_['layer2']==True:
            arquivo2.write('Pesos - camada oculta 2')
            arquivo2.write('\n')
            arquivo2.write(str(k[2]))
            arquivo2.write('\n')
            arquivo2.write('Bias - camada oculta 2')
            arquivo2.write('\n')
            arquivo2.write(str(k[3]))
            arquivo2.write('\n')
            arquivo2.write('Pesos - camada saída')
            arquivo2.write('\n')
            arquivo2.write(str(k[4].transpose()))
            arquivo2.write('\n')
        
        arquivo2.close()
        
        c=document.add_heading('RN')
        c.alignment=1
    
        document.add_paragraph('Número de neurônios: {}'.format(int(random_result.best_params_['n'])))
        document.add_paragraph('Taxa de aprendizado: {:e}'.format(random_result.best_params_['ln']))
        document.add_paragraph('Número de épocas: {}'.format(int(random_result.best_params_['epochs'])))
        document.add_paragraph('2° camada: {}'.format(random_result.best_params_['layer2']))
        document.add_paragraph('Função de ativação: {}'.format(random_result.best_params_['fun']))
        
        a=document.add_heading('Erros')
        a.alignment=1
        p=document.add_paragraph()
        p.add_run('Dados de teste').bold = True
        document.add_paragraph('Erro relativo médio: {}'.format(self.dici_erros['errorel_test']),style='List Bullet')
        document.add_paragraph('Coeficiente de correlação: {}'.format(self.dici_erros['cor_test']),style='List Bullet')
        document.add_paragraph('Coeficiente de determinação: {}'.format(self.dici_erros['det_test']),style='List Bullet')
        document.add_paragraph('MSE: {}'.format(self.dici_erros['mse_test']),style='List Bullet')
        document.add_paragraph('RMSE: {}'.format(self.dici_erros['rmse_test']),style='List Bullet')
        a=document.add_picture('Arquivos/RN/Gráficos/'+self.name+'_teste.png')
        a.alignment=1
    
        p=document.add_paragraph()
        p.add_run('Dados de treino').bold = True
        document.add_paragraph('Erro relativo médio: {}'.format(self.dici_erros['errorel_train']),style='List Bullet')
        document.add_paragraph('Coeficiente de correlação: {}'.format(self.dici_erros['cor_train']),style='List Bullet')
        document.add_paragraph('Coeficiente de determinação: {}'.format(self.dici_erros['det_train']),style='List Bullet')
        document.add_paragraph('MSE: {}'.format(self.dici_erros['mse_train']),style='List Bullet')
        document.add_paragraph('RMSE: {}'.format(self.dici_erros['rmse_train']),style='List Bullet')
        a=document.add_picture('Arquivos/RN/Gráficos/'+self.name+'_treino.png')
        a.alignment=1

        a=document.add_heading('Pesos')
        a.alignment=1
      
        k=self.model.get_weights()
        document.add_paragraph('Pesos - camada oculta 1')
        document.add_paragraph(str(k[0]))
        document.add_paragraph('Bias - camada oculta')
        document.add_paragraph(str(k[1]))
        if random_result.best_params_['layer2']==False:
            document.add_paragraph('Pesos - camada saída')
            document.add_paragraph(str(k[2].transpose()))

        elif random_result.best_params_['layer2']==True:
            document.add_paragraph('Pesos - camada oculta 2')
            document.add_paragraph(str(k[2]))
            document.add_paragraph('Bias - camada oculta 2')
            document.add_paragraph(str(k[3]))
            document.add_paragraph('Pesos - camada saída')
            document.add_paragraph(str(k[4].transpose()))
    
        
        a=document.add_heading('Iterações')
        a.alignment=1
   
        tabela(self.iteracao,document)
        document.add_page_break()
        
        
        ################## Regressões  ################
        
        degree=0
        for reg in ['RL','RP2','RP3','RP4']:
            degree+=1
            self.par=PolynomialFeatures(degree)
            self.line=LinearRegression()
            
            self.poly= make_pipeline(self.par, self.line,verbose=1)
            self.poly.fit(self.x_train,self.y_train)
       

            joblib.dump(self.poly, 'Arquivos/'+reg+'/Modelos/'+self.name+'_modelo.mdl')
            
            
            arquivo3=open('Arquivos/'+reg+'/Coeficientes/'+self.name+'_coeficientes.txt','w')
            arquivo3.write(str(self.line.coef_))
            arquivo3.close()
            
            pred_train0=self.poly.predict(self.x_train)
            pred_train0_r=norm_rev_y(pred_train0,self.train_stats_y)
            pred_train=np.array([])
            for item in pred_train0_r:
                pred_train=np.append(pred_train,round(float(item),2))
            
            pred_test0 = self.poly.predict(self.x_test)
            pred_test0_r=norm_rev_y(pred_test0,self.train_stats_y)
            pred_test=np.array([])
            for item in pred_test0_r:
                pred_test=np.append(pred_test,round(float(item),2))
            
            poui=(self.poly.score(self.x_test,self.y_test))
            
            cor_train_full=np.corrcoef(self.y_train_r,pred_train)[1,0]
            cor_test_full=np.corrcoef(self.y_test_r,pred_test)[1,0]
        
            det_test = np.round(cor_test_full**2,2)
            det_train = np.round(cor_train_full**2,2)
            
            erro(self.y_test_r,pred_test,self.y_train_r,pred_train,reg,self.name)
            graficos_ind(pred_test,self.y_test_r,pred_train,
                         self.y_train_r,self.u_saida,reg ,self.name,
                         reg+' $R^2$={:.2f}'.format(det_test),reg+' $R^2$={:.2f}'.format(det_train))
            

            
            ax.scatter(self.y_test_r,pred_test,label=reg+' $R^2$={:.2f}'.format(det_test))
            ax2.scatter(self.y_train_r,pred_train,label=reg+' $R^2$={:.2f}'.format(det_train))
        
            c=document.add_heading(reg)
            c.alignment=1
            
            c=document.add_heading('Coeficientes')
            c.alignment=1
            document.add_paragraph(str(self.line.coef_))
 
            a=document.add_heading('Erros')
            a.alignment=1
            p=document.add_paragraph()
            p.add_run('Dados de teste').bold = True
            document.add_paragraph('Erro relativo médio: {}'.format(self.dici_erros['errorel_test']),style='List Bullet')
            document.add_paragraph('Coeficiente de correlação: {}'.format(self.dici_erros['cor_test']),style='List Bullet')
            document.add_paragraph('Coeficiente de determinação: {}'.format(self.dici_erros['det_test']),style='List Bullet')
            document.add_paragraph('MSE: {}'.format(self.dici_erros['mse_test']),style='List Bullet')
            document.add_paragraph('RMSE: {}'.format(self.dici_erros['rmse_test']),style='List Bullet')
            a=document.add_picture('Arquivos/'+reg+'/Gráficos/'+self.name+'_teste.png')
            a.alignment=1
        
            p=document.add_paragraph()
            p.add_run('Dados de treino').bold = True
            document.add_paragraph('Erro relativo médio: {}'.format(self.dici_erros['errorel_train']),style='List Bullet')
            document.add_paragraph('Coeficiente de correlação: {}'.format(self.dici_erros['cor_train']),style='List Bullet')
            document.add_paragraph('Coeficiente de determinação: {}'.format(self.dici_erros['det_train']),style='List Bullet')
            document.add_paragraph('MSE: {}'.format(self.dici_erros['mse_train']),style='List Bullet')
            document.add_paragraph('RMSE: {}'.format(self.dici_erros['rmse_train']),style='List Bullet')
            a=document.add_picture('Arquivos/'+reg+'/Gráficos/'+self.name+'_treino.png')
            a.alignment=1
            document.add_page_break()
        
        ax.legend()
        menor=min([ax.get_xlim()[0],ax.get_ylim()[0]])
        maior=max([ax.get_xlim()[1],ax.get_ylim()[1]])
        ax.set_xlim(menor,maior)
        ax.set_ylim(menor,maior)
        ax.plot([menor,maior], [menor,maior],c='black',linewidth=0.65)
        ax.title.set_text('Modelos de previsão (dados de teste)')
        ax.set(xlabel="Valor real ({})".format(self.u_saida),ylabel="Previsão ({})".format(self.u_saida))
        fig.savefig('Arquivos/Gráficos/'+self.name+'_teste.png',bbox_inches='tight')
        
        ax2.legend()
        menor=min([ax2.get_xlim()[0],ax2.get_ylim()[0]])
        maior=max([ax2.get_xlim()[1],ax2.get_ylim()[1]])
        ax2.set_xlim(menor,maior)
        ax2.set_ylim(menor,maior)
        ax2.plot([menor,maior], [menor,maior],c='black',linewidth=0.65)
        ax2.title.set_text('Modelos de previsão (dados de treino)')
        ax2.set(xlabel="Valor real ({})".format(self.u_saida),ylabel="Previsão ({})".format(self.u_saida))
        fig2.savefig('Arquivos/Gráficos/'+self.name+'_treino.png',bbox_inches='tight')
       
        plt.close
        
        data_train={'Valor real':self.y_train_r}
        data_train.update(self.data_train_erro)
        data_test={'Valor real':self.y_test_r}
        data_test.update(self.data_test_erro)
        self.data_train=np.round(pd.DataFrame(data=data_train),2)
        self.data_test=np.round(pd.DataFrame(data=data_test),2)
        
        a=document.add_heading('Geral')
        a.alignment=1
        a=document.add_picture('Arquivos/Gráficos/'+self.name+'_teste.png')
        a.alignment=1
        a=document.add_picture('Arquivos/Gráficos/'+self.name+'_treino.png')
        a.alignment=1
        document.add_page_break()
        p=document.add_paragraph()
        p.add_run('Dados de teste').bold = True
        tabela(self.data_test,document)
        document.add_paragraph()
        p=document.add_paragraph()
        p.add_run('Dados de treino').bold = True
        tabela(self.data_train,document)
        
        document.save('Arquivos/Relatórios/'+self.name+'_relatorio.docx')
        
    def variaveis_exibir(self):
        
        self.errorel_test={}
        self.errorel_train={}
        self.cor_test={}
        self.cor_train={}
        self.det_test={}
        self.det_train={}
        self.mse_test={}
        self.mse_train={}
        self.rmse_test={}
        self.rmse_train={}
    
        
        with open('Arquivos/RN/Parâmetros/'+self.name+'_parametros.json', 'r') as json_file:
            parametros = json.load(json_file)   
        self.neuronios=parametros['n']
        self.ln=parametros['ln']
        self.epocas=parametros['epochs']    
        self.funcao=parametros['fun']
        self.layer=parametros['layer2']
        if self.layer==True:
            self.layer=2
        elif self.layer==False:
            self.layer=1
        
        for reg in ['RL','RP2','RP3','RP4','RN']:
        
            with open('Arquivos/'+reg+'/Erros/'+self.name+'_erros.json', 'r') as json_file:
                rn_erros = json.load(json_file)
            self.errorel_test[reg]=rn_erros['errorel_test']
            self.errorel_train[reg]=rn_erros['errorel_train']
            self.cor_test[reg]=rn_erros['cor_test']
            self.cor_train[reg]=rn_erros['cor_train']
            self.det_test[reg]=rn_erros['det_test']
            self.det_train[reg]=rn_erros['det_train']
            self.mse_test[reg]=rn_erros['mse_test']
            self.mse_train[reg]=rn_erros['mse_train']
            self.rmse_test[reg]=rn_erros['rmse_test']
            self.rmse_train[reg]=rn_erros['rmse_train']
         
    def carregar(self):
        from tensorflow.keras.models import load_model
        
        self.model={}
        self.model['RN']=load_model('Arquivos/RN/Modelos/'+self.name+'_modelo.mdl')
        
        for reg in ['RL','RP2','RP3','RP4']:
            self.model[reg]=joblib.load('Arquivos/'+reg+'/Modelos/'+self.name+'_modelo.mdl')

    def pred(self,v,f,a):
        self.pred={}
        r=[[v,f,a]]
        for reg in ['RL','RP2','RP3','RP4','RN']:
            self.r=norm_x(r,self.train_stats_x)
            self.pred[reg]=self.model[reg].predict(self.r)
            self.pred[reg]=norm_rev_y(self.pred[reg],self.train_stats_y)
            self.pred[reg]=np.round(float(self.pred[reg]),2)
        return(self.pred)